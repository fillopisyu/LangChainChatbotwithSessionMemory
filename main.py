import docx
from langchain import hub  # LangChain promptlarını çekmek için kullanılır
from langchain_chroma import Chroma  # Chroma vektör veritabanını kullanmak için
from langchain_core.output_parsers import StrOutputParser  # Çıktıyı string formatında almak için
from langchain_core.runnables import RunnablePassthrough  # Veriyi doğrudan iletmek için kullanılır
from langchain_openai import OpenAIEmbeddings  # OpenAI'nin gömülü (embedding) modelini kullanmak için
from langchain_text_splitters import RecursiveCharacterTextSplitter  # Metin parçalamak için
from dotenv import load_dotenv  # Çevresel değişkenleri yüklemek için
from langchain_openai import ChatOpenAI  # OpenAI'nin dil modeliyle çalışmak için
from langchain.schema import Document  # Document sınıfını içe aktar


# Word dosyasını okuma ve metni çıkarma işlevi
def read_word_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

# Çevresel değişkenleri (.env dosyasını) yükle
load_dotenv()

# OpenAI'nin GPT-3.5-turbo modelini kullanmak için bir ChatOpenAI nesnesi oluştur
llm = ChatOpenAI(model="gpt-3.5-turbo")

# Word dosyasındaki metni oku
word_file_path = r"C:\Users\ismsa\OneDrive\Masaüstü\Kartlar.docx"  # Word dosyanızın yolu
doc_content = read_word_file(word_file_path)

# Metni belirli boyutlarda parçalara ayırmak için bir text splitter oluştur
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

# Word dosyasından alınan içerikleri belgeler listesine ekleyelim
docs = [Document(page_content=doc_content)]

# Belgeleri parçalara ayır
splits = text_splitter.split_documents(docs)

# Metni 'Document' formatında parçalara ayır (Her split bir Document nesnesi olacak)
documents = [Document(page_content=chunk.page_content) for chunk in splits]

# Chroma vektör veritabanı oluştur ve OpenAI'nin embedding modelini kullanarak metinleri indeksle
vectorstore = Chroma.from_documents(documents=documents, embedding=OpenAIEmbeddings())

# Vektör veritabanından ilgili belgeleri getirmek için bir retriever oluştur
retriever = vectorstore.as_retriever()

# LangChain hub'dan RAG (Retrieval-Augmented Generation) promptunu çek
prompt = hub.pull("rlm/rag-prompt")

# Belgeleri formatlamak için bir fonksiyon tanımla
def format_docs(docs):
    # Belgelerin içeriklerini birleştirerek tek bir string döndür
    return "\n\n".join(doc.page_content for doc in docs)

# RAG (Retrieval-Augmented Generation) zinciri oluştur
rag_chain = (
    {"context": retriever | format_docs, "question": RunnablePassthrough()}  # Soruyu doğrudan geçir
    | prompt  # Prompt'u uygulama
    | llm  # Dil modelinden cevap al
    | StrOutputParser()  # Cevabı string formatına çevir
)

# Ana çalışma bloğu
if __name__ == "__main__":
    # Boş bir soru göndererek Word dosyasının içeriği ile işlem yap
    for chunk in rag_chain.stream("İnternet Bankacılığı'nda VakıfPay Ödemesi Nasıl Yaparım?    "):
        # Cevabı parça parça yazdır
        print(chunk, end="", flush=True)
